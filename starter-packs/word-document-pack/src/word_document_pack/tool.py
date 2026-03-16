"""Create Word documents with headings, paragraphs, tables, and lists."""

from __future__ import annotations

import os
import tempfile

from docx import Document
from docx.shared import Pt


def _count_words(doc: Document) -> int:
    """Count words across all paragraphs and table cells."""
    count = 0
    for para in doc.paragraphs:
        count += len(para.text.split())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                count += len(cell.text.split())
    return count


def run(
    content: list[dict],
    output_path: str = "",
    title: str = "",
) -> dict:
    """Create a Word document from structured content blocks.

    Args:
        content: List of content blocks. Each block is a dict with:
            - type: "heading", "paragraph", "table", or "list"
            - text: The text content (for tables, rows separated by newlines,
                    columns separated by '|'; for lists, items separated by newlines)
            - level: Heading level (1-9) or list nesting level (default 0)
        output_path: Where to save the .docx file. Auto-generated if empty.
        title: Optional document title added as a top-level heading.

    Returns:
        Dictionary with output_path and word_count.
    """
    if not content:
        raise ValueError("At least one content block is required.")

    doc = Document()

    # Add title if provided
    if title:
        doc.add_heading(title, level=0)

    for block in content:
        block_type = block.get("type", "paragraph")
        text = block.get("text", "")
        level = block.get("level", 1)

        if block_type == "heading":
            heading_level = max(1, min(int(level), 9))
            doc.add_heading(text, level=heading_level)

        elif block_type == "paragraph":
            doc.add_paragraph(text)

        elif block_type == "table":
            # Parse table from text: rows separated by newlines, columns by '|'
            rows = [
                [cell.strip() for cell in row.split("|")]
                for row in text.strip().split("\n")
                if row.strip()
            ]
            if not rows:
                continue

            num_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = "Table Grid"

            for ri, row_data in enumerate(rows):
                for ci, cell_text in enumerate(row_data):
                    if ci < num_cols:
                        table.cell(ri, ci).text = cell_text

        elif block_type == "list":
            items = [item.strip() for item in text.split("\n") if item.strip()]
            list_level = max(0, int(level)) if level else 0
            for item in items:
                paragraph = doc.add_paragraph(item, style="List Bullet")
                paragraph.paragraph_format.left_indent = Pt(18 * list_level)

        else:
            # Unknown type; treat as paragraph
            doc.add_paragraph(text)

    # Determine output path
    if not output_path:
        safe_title = "".join(
            c if c.isalnum() or c in (" ", "-", "_") else "_" for c in (title or "document")
        ).strip()
        output_path = os.path.join(
            tempfile.gettempdir(), f"{safe_title}.docx"
        )

    output_path = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    word_count = _count_words(doc)

    return {
        "output_path": output_path,
        "word_count": word_count,
    }
